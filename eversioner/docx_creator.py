from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from typing import List, Dict
import os
from datetime import datetime

def salvar_prova(prova: List[Dict], gabarito: List[str], versao: int, avc_numero: int, data_prova: str):
    """
    Salva uma versão da prova em um arquivo .docx.

    Args:
        prova: Lista de questões da prova.
        gabarito: Lista com o gabarito da prova.
        versao: Número da versão da prova.
        avc_numero: Número da AVC.
        data_prova: Data da prova.
    """
    doc = Document()
    
    # Configurar página A4 com margens médias
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Cm(2.54)
    
    # Configurar estilo padrão
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(12 * 1)  # 1.0 vezes o tamanho da fonte
    
    # 1. Adicionar logo
    doc.add_picture('logo_unisa.png', width=Cm(1.5*16/9), height=Cm(1.5))  # Mantém a proporção 16:9
    
    # 2. Adicionar cabeçalho
    for texto in [
        "UNIVERSIDADE DE SANTO AMARO - Curso de MEDICINA",
        f"NÚCLEO DE SAÚDE DA MULHER - AVC {avc_numero}, Versão {versao}",
        f"São Paulo, {data_prova}"
    ]:
        p = doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(14)
    
    doc.add_paragraph()  # Pular um parágrafo
    
    # 3. Adicionar campos para nome, assinatura e RA
    for texto in [
        "NOME: ___________________________________________________",
        "ASSINATURA: ______________________ RA: __________________"
    ]:
        p = doc.add_paragraph(texto)
        for run in p.runs:
            run.font.size = Pt(12)
    
    # 4. Adicionar 2 parágrafos antes das questões
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Adicionar questões
    for questao in prova:
        p = doc.add_paragraph(f"{questao['numero']}. {questao['enunciado']}")
        p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for i, alternativa in enumerate(questao['alternativas']):
            p = doc.add_paragraph(f"    {chr(65 + i)}) {alternativa}")
        doc.add_paragraph()  # Espaço entre questões
    
    # 5. Adicionar quebra de página e repetir cabeçalho
    doc.add_page_break()
    
    # Repetir itens 1, 2 e 3
    doc.add_picture('logo_unisa.png', width=Cm(2*16/9), height=Cm(1.75))
    
    for texto in [
        "UNIVERSIDADE DE SANTO AMARO - Curso de MEDICINA",
        f"NÚCLEO DE SAÚDE DA MULHER - AVC {avc_numero}, Versão {versao}",
        f"São Paulo, {data_prova}"
    ]:
        p = doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    for texto in [
        "NOME: ___________________________________________________",
        "ASSINATURA: ______________________ RA: __________________"
    ]:
        p = doc.add_paragraph(texto)
        for run in p.runs:
            run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()

    # Adicionar título "RESPOSTAS"
    p = doc.add_paragraph("RESPOSTAS")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    
    # 6. Adicionar parágrafo de instruções
    p = doc.add_paragraph("ATENÇÃO: Apenas as respostas assinaladas nesta tabela serão consideradas. Não rasure. Respostas rasuradas, anotações ou marcações nas questões serão desconsideradas.")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].font.size = Pt(12)
    
    # 7. Adicionar tabela de respostas
    table = doc.add_table(rows=len(prova)+1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Definir largura da tabela
    table.width = Cm(4)
    
    # Cabeçalho da tabela
    header_cells = table.rows[0].cells
    header_cells[0].text = "Questão"
    header_cells[1].text = "Resposta"
    
    # Preencher a tabela
    for i, questao in enumerate(prova, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = str(i)
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Ajustar o tamanho das células
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(14)
    
    # Salvando o arquivo
    if not os.path.exists('provas'):
        os.makedirs('provas')
    doc.save(f'provas/Prova_Versao_{versao}.docx')
    print(f"Prova versão {versao} salva com sucesso.")


def salvar_gabaritos(gabaritos: List[List[str]], avc_numero: int, data_prova: str):
    """
    Salva todos os gabaritos em um único arquivo .docx.

    Args:
        gabaritos: Lista de gabaritos de todas as versões.
        avc_numero: Número da AVC.
        data_prova: Data da prova.
    """
    doc = Document()
    
    # Configurar página A4 com margens médias
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Cm(2.54)
    
    # Configurar estilo padrão
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    style.font.name = 'Arial'
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(12 * 1.0)  # 1.0 vezes o tamanho da fonte
    
    # 1. Adicionar logo
    doc.add_picture('logo_unisa.png', width=Cm(1.5*16/9), height=Cm(1.5))  # Mantém a proporção 16:9
    
    # 2. Adicionar cabeçalho
    for texto in [
        "UNIVERSIDADE DE SANTO AMARO -- Curso de MEDICINA",
        f"NÚCLEO DE SAÚDE DA MULHER - AVC {avc_numero}",
        f"SÃO PAULO, {data_prova}"
    ]:
        p = doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(14)
    
    # Pular 2 parágrafos
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Adicionar título "GABARITOS"
    titulo = doc.add_paragraph("GABARITOS")
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].bold = True
    titulo.runs[0].font.size = Pt(14)

    # Criar tabela de gabaritos
    num_questoes = len(gabaritos[0])
    num_versoes = len(gabaritos)
    table = doc.add_table(rows=num_questoes+1, cols=num_versoes+1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Preencher cabeçalho da tabela
    header_cells = table.rows[0].cells
    header_cells[0].text = "Questão"
    for i in range(num_versoes):
        header_cells[i+1].text = f"Versão {i+1}"

    # Preencher a tabela com os gabaritos
    for i in range(num_questoes):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = str(i+1)
        for j in range(num_versoes):
            row_cells[j+1].text = gabaritos[j][i].split('-')[1]  # Pega apenas a letra da resposta

    # Formatar a tabela
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(14)

    # Salvando o arquivo
    if not os.path.exists('provas'):
        os.makedirs('provas')
    doc.save('provas/Gabaritos.docx')
    print("Gabaritos salvos com sucesso.")